import pdfplumber
from openai_helper import summarize_text, generate_improvement_suggestions
from docx import Document
import os
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path):
    logger.info(f"Extracting text from PDF: {pdf_path}")
    text_by_page = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                text_by_page.append(text)
        logger.info(f"Successfully extracted text from {len(text_by_page)} pages")
        return text_by_page
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def create_summary_docx(slide_summaries, improvement_suggestions, output_path):
    logger.info(f"Creating summary DOCX at: {output_path}")
    doc = Document()
    doc.add_heading('Presentation Summary and Analysis', 0)

    for i, (summary, suggestions) in enumerate(zip(slide_summaries, improvement_suggestions), 1):
        doc.add_heading(f'Slide {i}', level=1)
        doc.add_paragraph(f'Summary: {summary}')
        doc.add_paragraph(f'Improvement Suggestions: {suggestions}')

    doc.save(output_path)
    logger.info(f"Summary saved to {output_path}")
    return output_path

def analyze_pdf_presentation(pdf_path):
    logger.info(f"Analyzing PDF presentation: {pdf_path}")
    try:
        # Step 1: Extract text from PDF
        slides_text = extract_text_from_pdf(pdf_path)

        # Step 2: Summarize each slide and generate suggestions
        logger.info("Generating summaries and improvement suggestions")
        slide_summaries = [summarize_text(slide) for slide in slides_text]
        improvement_suggestions = [generate_improvement_suggestions(slide) for slide in slides_text]

        # Step 3: Create a DOCX report
        output_folder = os.path.dirname(pdf_path)
        output_path = os.path.join(output_folder, "summary.docx")
        return create_summary_docx(slide_summaries, improvement_suggestions, output_path)
    except Exception as e:
        logger.error(f"Error analyzing PDF presentation: {str(e)}")
        raise